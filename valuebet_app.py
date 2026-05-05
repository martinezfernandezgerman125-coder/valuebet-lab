import streamlit as st
import pandas as pd
import numpy as np
from datetime import datetime, date
import json
import sqlite3
import requests
from scipy.stats import poisson

st.set_page_config(page_title="ValueBet Lab", page_icon="🔍", layout="wide")

# ============================================================
# CONFIGURACIÓN DE IA
# ============================================================

GROQ_API_KEY = st.secrets.get("GROQ_API_KEY", "")
IA_DISPONIBLE = len(GROQ_API_KEY) > 10

st.markdown("""
<style>
    .stApp { background-color: #0A1F14; }
    h1, h2, h3 { color: #5DCAA5 !important; }
    .stButton > button {
        background-color: #0F6E56; color: white; border: none;
        border-radius: 8px; padding: 10px 24px; font-weight: 600;
    }
    .stButton > button:hover { background-color: #1D9E75; }
    .card {
        background: linear-gradient(135deg, #0A1F14, #0F6E56);
        border: 1px solid #1D9E75; border-radius: 12px;
        padding: 16px; margin: 8px 0;
    }
    .badge-ia {
        background: #1D9E75; padding: 2px 10px; border-radius: 10px;
        font-size: 11px; color: white; margin-left: 10px;
    }
</style>
""", unsafe_allow_html=True)

# ============================================================
# BASE DE DATOS
# ============================================================

def init_db():
    conn = sqlite3.connect("valuebet.db")
    c = conn.cursor()
    c.execute("""CREATE TABLE IF NOT EXISTS encuentros (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        fecha DATE NOT NULL,
        equipo_local TEXT NOT NULL,
        equipo_visitante TEXT NOT NULL,
        estado TEXT DEFAULT 'pendiente',
        resultado_analisis TEXT,
        doc_text TEXT,
        fecha_creacion TIMESTAMP DEFAULT CURRENT_TIMESTAMP
    )""")
    conn.commit()
    conn.close()

init_db()

def get_db():
    conn = sqlite3.connect("valuebet.db")
    conn.row_factory = sqlite3.Row
    return conn

def add_match(fecha, local, visitante):
    conn = get_db()
    conn.execute("INSERT INTO encuentros (fecha, equipo_local, equipo_visitante) VALUES (?, ?, ?)", (fecha, local, visitante))
    conn.commit()
    id_ = conn.execute("SELECT last_insert_rowid()").fetchone()[0]
    conn.close()
    return id_

def get_matches(fecha):
    conn = get_db()
    m = conn.execute("SELECT * FROM encuentros WHERE fecha = ? ORDER BY equipo_local", (fecha,)).fetchall()
    conn.close()
    return m

def get_match(id_):
    conn = get_db()
    m = conn.execute("SELECT * FROM encuentros WHERE id = ?", (id_,)).fetchone()
    conn.close()
    return m

def update_match(id_, field, value):
    conn = get_db()
    conn.execute(f"UPDATE encuentros SET {field}=? WHERE id=?", (value, id_))
    conn.commit()
    conn.close()

def delete_match(id_):
    conn = get_db()
    conn.execute("DELETE FROM encuentros WHERE id = ?", (id_,))
    conn.commit()
    conn.close()

# ============================================================
# ANÁLISIS CON IA (Groq + Llama 3)
# ============================================================

def analyze_with_ia(local, visitante, doc_text=""):
    """
    Analiza el partido usando Llama 3 70B via Groq (GRATIS).
    Si falla, usa Poisson local como respaldo.
    """
    
    datos_extra = ""
    if doc_text:
        datos_extra = f"\nDATOS ADICIONALES DEL DOCUMENTO:\n{doc_text[:2000]}"
    
    prompt = f"""Eres un analista de datos deportivos experto en value betting.
Analiza este partido de fútbol y genera un análisis detallado.

PARTIDO: {local} vs {visitante}{datos_extra}

INSTRUCCIONES:
1. Estima los goles esperados para cada equipo (basado en estadísticas reales si conoces los equipos)
2. Calcula probabilidades de resultado
3. Identifica 3-5 picks value con cuotas realistas (1.50-1.80)
4. Para cada pick, da una justificación basada en datos

RESPONDE SOLO CON JSON (sin markdown, sin explicación extra):
{{
    "goles_local": 1.5,
    "goles_visitante": 0.8,
    "prob_local": 55.0,
    "prob_empate": 25.0,
    "prob_visitante": 20.0,
    "prob_over_2_5": 60.0,
    "prob_btts": 65.0,
    "picks": [
        {{
            "mercado": "Doble Oportunidad",
            "pick": "1X",
            "probabilidad": 80.0,
            "cuota": 1.57,
            "justificacion": "Razonamiento basado en datos del equipo"
        }}
    ]
}}"""
    
    try:
        response = requests.post(
            "https://api.groq.com/openai/v1/chat/completions",
            headers={
                "Authorization": f"Bearer {GROQ_API_KEY}",
                "Content-Type": "application/json"
            },
            json={
                "model": "llama3-70b-8192",
                "messages": [
                    {"role": "system", "content": "Eres un analista deportivo experto. Responde solo JSON valido."},
                    {"role": "user", "content": prompt}
                ],
                "temperature": 0.1,
                "max_tokens": 2000
            },
            timeout=30
        )
        
        if response.status_code == 200:
            content = response.json()["choices"][0]["message"]["content"]
            content = content.strip().replace("```json", "").replace("```", "")
            resultado = json.loads(content)
            
            # Asegurar que tiene todos los campos necesarios
            if "goles_local" not in resultado:
                resultado["goles_local"] = 1.4
            if "prob_over_2_5" not in resultado:
                resultado["prob_over_2_5"] = 55.0
            if "prob_btts" not in resultado:
                resultado["prob_btts"] = 55.0
            if "picks" not in resultado:
                resultado["picks"] = []
            
            resultado["modelo"] = "Llama 3 70B (Groq)"
            return resultado
            
    except Exception as e:
        pass  # Si falla, usar Poisson
    
    # Fallback: Poisson
    return analyze_poisson(local, visitante)

def analyze_poisson(local, visitante):
    """Modelo Poisson de respaldo"""
    lambda_h = 1.4
    lambda_a = 1.1
    
    # Ajustes por equipos conocidos
    known = {
        "real madrid": (2.5, 0.5), "barcelona": (2.3, 0.6),
        "man city": (2.4, 0.5), "liverpool": (2.1, 0.7),
        "levante": (0.9, 1.8), "elche": (0.8, 1.9),
        "alaves": (0.8, 1.7), "cadiz": (0.7, 1.8)
    }
    
    l_lower = local.lower()
    v_lower = visitante.lower()
    
    for nombre, (ataque, defensa) in known.items():
        if nombre in l_lower:
            lambda_h = ataque + 0.3
        if nombre in v_lower:
            lambda_a = defensa
    
    max_g = 6
    probs = np.zeros((max_g+1, max_g+1))
    for i in range(max_g+1):
        for j in range(max_g+1):
            probs[i][j] = poisson.pmf(i, lambda_h) * poisson.pmf(j, lambda_a)
    probs = probs / probs.sum()
    
    p_local = float(np.sum(probs * np.tri(max_g+1, max_g+1, -1).T))
    p_empate = float(np.trace(probs))
    p_visitante = float(np.sum(probs * np.tri(max_g+1, max_g+1, -1)))
    p_over = float(1 - (probs[0,0]+probs[0,1]+probs[1,0]+probs[1,1]))
    p_btts = float(np.sum(probs[1:, 1:]))
    
    picks = []
    if p_local + p_empate > 0.70:
        cuota = round(1/(p_local+p_empate+0.05), 2)
        picks.append({"mercado": "Doble Oportunidad", "pick": "1X", "probabilidad": round((p_local+p_empate)*100, 1), "cuota": cuota, "justificacion": f"Poisson: {round((p_local+p_empate)*100)}% probabilidad"})
    if p_over > 0.45:
        cuota = round(1/(p_over+0.08), 2)
        picks.append({"mercado": "Goles", "pick": "Over 2.5", "probabilidad": round(p_over*100, 1), "cuota": cuota, "justificacion": f"Goles esperados: {round(lambda_h+lambda_a, 1)}"})
    if p_btts > 0.45:
        cuota = round(1/(p_btts+0.1), 2)
        picks.append({"mercado": "BTTS", "pick": "Si", "probabilidad": round(p_btts*100, 1), "cuota": cuota, "justificacion": f"BTTS: {round(p_btts*100)}%"})
    if p_local > 0.40:
        cuota = round(1/(p_local+0.1), 2)
        picks.append({"mercado": "Resultado", "pick": f"Gana {local}", "probabilidad": round(p_local*100, 1), "cuota": cuota, "justificacion": f"Victoria local: {round(p_local*100)}%"})
    if p_over > 0.60:
        cuota = round(1/(p_over+0.05), 2)
        picks.append({"mercado": "Intervalo", "pick": "1-3 goles", "probabilidad": round(p_over*85, 1), "cuota": cuota, "justificacion": "Mercado estrella para este tipo de partidos"})
    
    return {"goles_local": round(lambda_h, 2), "goles_visitante": round(lambda_a, 2), "prob_local": round(p_local*100, 1), "prob_empate": round(p_empate*100, 1), "prob_visitante": round(p_visitante*100, 1), "prob_over_2_5": round(p_over*100, 1), "prob_btts": round(p_btts*100, 1), "picks": picks, "modelo": "Poisson Ajustado"}

# ============================================================
# SUBIR DOCUMENTOS
# ============================================================

def extract_text(uploaded_file):
    """Extrae texto de archivos subidos"""
    try:
        from docx import Document
        from io import BytesIO
        doc = Document(BytesIO(uploaded_file.getvalue()))
        text = "\n".join([p.text for p in doc.paragraphs])
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += "\n" + cell.text
        return text[:10000]
    except:
        try:
            return uploaded_file.getvalue().decode("utf-8", errors="ignore")[:10000]
        except:
            return ""

# ============================================================
# INTERFAZ DE USUARIO
# ============================================================

def main():
    # Header
    ia_badge = '<span class="badge-ia">🤖 IA ACTIVA</span>' if IA_DISPONIBLE else '<span class="badge-ia" style="background:#666;">📊 MODO LOCAL</span>'
    st.markdown(f"<h1 style='text-align:center;'>🔍 ValueBet Lab {ia_badge}</h1>", unsafe_allow_html=True)
    st.markdown("<p style='text-align:center;color:#9FE1CB;'>Analizador de Partidos con IA</p>", unsafe_allow_html=True)
    
    # Sidebar
    with st.sidebar:
        st.markdown("### 📅 Calendario")
        fecha = st.date_input("Seleccionar fecha", datetime.now(), 
                             min_value=date(2024,1,1), max_value=date(2030,12,31),
                             key="fecha_selector")
        
        st.markdown("---")
        st.markdown("### ➕ Nuevo Partido")
        local = st.text_input("🏠 Equipo Local")
        visit = st.text_input("✈️ Equipo Visitante")
        
        col_btn1, col_btn2 = st.columns(2)
        with col_btn1:
            if st.button("➕ Añadir", use_container_width=True):
                if local and visit and local != visit:
                    add_match(fecha, local, visit)
                    st.success(f"✅ {local} vs {visit}")
                    st.rerun()
                else:
                    st.error("Completa los campos")
        
        with col_btn2:
            if st.button("🗑️ Limpiar", use_container_width=True):
                st.rerun()
        
        st.markdown("---")
        
        # Stats
        conn = get_db()
        total = conn.execute("SELECT COUNT(*) FROM encuentros").fetchone()[0]
        analizados = conn.execute("SELECT COUNT(*) FROM encuentros WHERE estado='analizado'").fetchone()[0]
        pendientes = total - analizados
        conn.close()
        
        st.markdown("### 📊 Estadísticas")
        col1, col2, col3 = st.columns(3)
        col1.metric("Total", total)
        col2.metric("Analizados", analizados)
        col3.metric("Pendientes", pendientes)
        
        if not IA_DISPONIBLE:
            st.markdown("---")
            st.warning("""
            💡 **Mejora tus análisis**
            
            Conecta Groq (IA gratuita) para análisis más precisos:
            1. Ve a https://console.groq.com
            2. Obtén tu API Key
            3. Crea `.streamlit/secrets.toml`
            """)
    
    # Main content
    st.markdown(f"### 📋 Partidos del {fecha.strftime('%d/%m/%Y')}")
    
    matches = get_matches(fecha)
    
    if not matches:
        st.info("ℹ️ No hay partidos para esta fecha. Añade uno en el panel izquierdo.")
    else:
        for i, m in enumerate(matches):
            estado = "🟢 Analizado" if m["estado"] == "analizado" else "🟡 Pendiente"
            
            with st.container():
                st.markdown(f"""
                <div class="card">
                    <div style="display:flex;justify-content:space-between;align-items:center;">
                        <div>
                            <strong style="color:#E1F5EE;font-size:18px;">{m['equipo_local']}</strong>
                            <span style="color:#9FE1CB;font-size:18px;"> vs </span>
                            <strong style="color:#E1F5EE;font-size:18px;">{m['equipo_visitante']}</strong>
                            <span style="color:#5DCAA5;margin-left:15px;">{estado}</span>
                        </div>
                    </div>
                </div>
                """, unsafe_allow_html=True)
                
                # Botones
                col_a, col_b, col_c, col_d = st.columns([2, 2, 2, 1])
                
                with col_a:
                    if st.button(f"📄 Subir DOC", key=f"doc_{m['id']}", use_container_width=True):
                        st.session_state[f"show_doc_{m['id']}"] = True
                
                with col_b:
                    if m["estado"] == "analizado":
                        if st.button(f"📊 Ver resultados", key=f"view_{m['id']}", use_container_width=True):
                            st.session_state[f"show_results_{m['id']}"] = not st.session_state.get(f"show_results_{m['id']}", False)
                    else:
                        if st.button(f"🔍 Analizar", key=f"go_{m['id']}", use_container_width=True):
                            with st.spinner("Analizando..."):
                                doc_text = m["doc_text"] if m["doc_text"] else ""
                                
                                if IA_DISPONIBLE:
                                    resultado = analyze_with_ia(m['equipo_local'], m['equipo_visitante'], doc_text)
                                else:
                                    resultado = analyze_poisson(m['equipo_local'], m['equipo_visitante'])
                                
                                update_match(m['id'], "resultado_analisis", json.dumps(resultado))
                                update_match(m['id'], "estado", "analizado")
                            st.success("✅ Análisis completado")
                            st.rerun()
                
                with col_c:
                    if st.button(f"🗑️ Eliminar", key=f"del_{m['id']}", use_container_width=True):
                        delete_match(m['id'])
                        st.rerun()
                
                with col_d:
                    if m['doc_text']:
                        st.markdown("📄✅")
                
                # Subir documento
                if st.session_state.get(f"show_doc_{m['id']}", False):
                    doc_file = st.file_uploader(f"Sube documento para {m['equipo_local']} vs {m['equipo_visitante']}", 
                                               type=['docx', 'txt', 'csv'], key=f"upload_{m['id']}")
                    if doc_file:
                        text = extract_text(doc_file)
                        update_match(m['id'], "doc_text", text)
                        st.success("✅ Documento guardado")
                        st.session_state[f"show_doc_{m['id']}"] = False
                        st.rerun()
                
                # Mostrar resultados
                if st.session_state.get(f"show_results_{m['id']}", False) and m["estado"] == "analizado":
                    r = json.loads(m['resultado_analisis'])
                    
                    col1, col2, col3 = st.columns(3)
                    col1.metric("⚽ Goles Local", r.get("goles_local", "?"))
                    col2.metric("📊 Total", round(r.get("goles_local", 0) + r.get("goles_visitante", 0), 1))
                    col3.metric("⚽ Goles Visit", r.get("goles_visitante", "?"))
                    
                    st.markdown(f"**Modelo:** {r.get('modelo', 'Poisson')}")
                    
                    st.markdown("### 📊 Probabilidades")
                    probs_df = pd.DataFrame({
                        "Resultado": [m['equipo_local'], "Empate", m['equipo_visitante']],
                        "Probabilidad": [f"{r.get('prob_local', 0)}%", f"{r.get('prob_empate', 0)}%", f"{r.get('prob_visitante', 0)}%"]
                    })
                    st.dataframe(probs_df, hide_index=True, use_container_width=True)
                    
                    col1, col2 = st.columns(2)
                    col1.metric("Over 2.5 Goles", f"{r.get('prob_over_2_5', 0)}%")
                    col2.metric("BTTS", f"{r.get('prob_btts', 0)}%")
                    
                    if r.get("picks"):
                        st.markdown("### 🏆 Picks Sugeridos")
                        for i, p in enumerate(r["picks"], 1):
                            st.markdown(f"""
                            <div class="card">
                                <h4 style="color:#5DCAA5;margin:0;">Pick #{i}</h4>
                                <p style="font-size:18px;margin:8px 0;">
                                    <strong>{p['mercado']}:</strong> {p['pick']}
                                </p>
                                <p>📊 Prob: {p['probabilidad']}% | 💰 Cuota: {p['cuota']}</p>
                                <p style="font-size:12px;color:#9FE1CB;">{p['justificacion']}</p>
                            </div>
                            """, unsafe_allow_html=True)

if __name__ == "__main__":
    main()
